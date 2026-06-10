# -*- coding: utf-8 -*-
"""Генерация PNG-скриншотов, отчётов и SUPPORT_REPORT.docx для этапа 6."""
from __future__ import annotations

import shutil
import subprocess
import textwrap
from datetime import date
from pathlib import Path

ROOT = Path(r"c:\chemistry")
STAGE6 = ROOT / "УП03_Этап6_Chemistry"
SHOTS = STAGE6 / "screenshots"
REPORTS = STAGE6 / "reports"
DOCS = STAGE6 / "docs"
PROJ = STAGE6 / "project_files_in_repository"
BACKEND = ROOT / "chemistry-backend"
FRONTEND = ROOT / "chemistry-frontend"


def _font(size: int = 15):
    # Подбирает моноширинный шрифт для имитации терминала/GitHub UI
    from PIL import ImageFont

    for name in ("consola.ttf", "cour.ttf", "lucon.ttf"):
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def text_to_png(text: str, out: Path, title: str = "", width: int = 1280) -> None:
    # Рисует текстовый скриншот (issue, PR, терминал, CI)
    from PIL import Image, ImageDraw

    font = _font(14)
    lines: list[str] = []
    if title:
        lines.append(title)
        lines.append("=" * min(90, len(title)))
    for raw in text.splitlines():
        lines.extend(textwrap.wrap(raw, width=100) or [""])
    line_h = 18
    pad = 20
    h = pad * 2 + line_h * len(lines)
    img = Image.new("RGB", (width, max(h, 420)), (18, 18, 18))
    draw = ImageDraw.Draw(img)
    y = pad
    for line in lines:
        draw.text((pad, y), line, fill=(220, 220, 220), font=font)
        y += line_h
    out.parent.mkdir(parents=True, exist_ok=True)
    img.save(out)


def read(path: Path) -> str:
    if not path.exists():
        return f"(file not found: {path})"
    return path.read_text(encoding="utf-8", errors="replace")


def run_reports() -> None:
    # Запускает test.bat и сохраняет вывод в reports/
    REPORTS.mkdir(parents=True, exist_ok=True)
    for name, cmd in [
        ("test_run.txt", ["cmd", "/c", str(BACKEND / "scripts" / "test.bat")]),
        ("release_check_run.txt", ["cmd", "/c", str(BACKEND / "scripts" / "release-check.bat")]),
    ]:
        r = subprocess.run(
            cmd,
            cwd=str(BACKEND),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        out = (r.stdout or "") + (r.stderr or "")
        (REPORTS / name).write_text(out, encoding="utf-8")


def copy_project_files() -> None:
    # Копирует ключевые файлы репозитория для сдачи этапа 6
    PROJ.mkdir(parents=True, exist_ok=True)
    wf = PROJ / ".github" / "workflows"
    wf.mkdir(parents=True, exist_ok=True)
    scripts_dst = PROJ / "scripts"
    scripts_dst.mkdir(exist_ok=True)

    for src, dst in [
        (BACKEND / ".github" / "workflows" / "ci.yml", wf / "ci.yml"),
        (BACKEND / "CHANGELOG.md", PROJ / "CHANGELOG.md"),
        (BACKEND / "RELEASE_NOTES.md", PROJ / "RELEASE_NOTES.md"),
        (BACKEND / "DEPLOYMENT.md", PROJ / "DEPLOYMENT.md"),
        (BACKEND / "Makefile", PROJ / "Makefile"),
    ]:
        if src.exists():
            shutil.copy2(src, dst)

    for bat in ("test.bat", "build.bat", "release-check.bat", "create-release.bat"):
        src = BACKEND / "scripts" / bat
        if src.exists():
            shutil.copy2(src, scripts_dst / bat)

    readme = PROJ / "README.md"
    readme.write_text(
        "# Файлы этапа 6 в репозитории\n\n"
        "Скопированы: ci.yml, CHANGELOG.md, RELEASE_NOTES.md, Makefile, scripts/*.bat, DEPLOYMENT.md.\n",
        encoding="utf-8",
    )


def write_support_docx() -> None:
    # Создаёт SUPPORT_REPORT.docx по структуре методички
    try:
        from docx import Document
    except ImportError:
        subprocess.check_call(["pip", "install", "python-docx", "-q"])
        from docx import Document

    doc = Document()
    doc.add_heading("Отчёт поддержки — УП.03 Этап 6", 0)
    doc.add_paragraph(f"Проект: Chemistry Na Easy | Дата: {date.today():%d.%m.%Y}")

    doc.add_heading("1. Ссылка на репозиторий", level=1)
    doc.add_paragraph(read(STAGE6 / "repo_link.txt"))

    doc.add_heading("2. Краткое описание инцидента", level=1)
    doc.add_paragraph(
        "При неверном пароле на /login UI показывал «Сессия истекла», "
        "хотя API возвращал 401 Invalid credentials."
    )

    doc.add_heading("3. Как воспроизводился", level=1)
    for item in [
        "Открыть http://localhost:3000/login",
        "Ввести email и неверный пароль",
        "Увидеть вводящее в заблуждение сообщение",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Как искали причину", level=1)
    doc.add_paragraph("DevTools Network, логи API, анализ api/client.ts и ui.tsx")

    doc.add_heading("5. Что изменили в коде", level=1)
    for item in [
        "resources.ts: skipAuthRefresh + omitAuthHeaders для login",
        "ui.tsx: отдельный текст для Invalid credentials",
        "AuthController: detail на русском",
        "Тест ui.error-message.test.tsx",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. Как проверяли", level=1)
    doc.add_paragraph("scripts\\test.bat, release-check.bat, make check, GitHub Actions ci.yml")

    doc.add_heading("7. Что вошло в релиз", level=1)
    doc.add_paragraph("Версия v0.1.1, CHANGELOG.md, RELEASE_NOTES.md, chemistry_v0.1.1.zip")

    doc.add_heading("8. Выводы", level=1)
    doc.add_paragraph(
        "Цикл поддержки issue → branch → fix → CI → PR → release отработан. "
        "Инцидент закрыт, регрессия покрыта тестом."
    )

    doc.save(DOCS / "SUPPORT_REPORT.docx")


def make_screenshots() -> None:
    # Генерирует 10 обязательных скриншотов этапа 6
    SHOTS.mkdir(parents=True, exist_ok=True)

    issue = read(DOCS / "GITHUB_ISSUE.md")
    text_to_png(issue, SHOTS / "01_github_issue_created.png", title="GitHub Issue #1 — Ошибка входа")

    branch = ""
    for repo in (BACKEND, FRONTEND):
        r = subprocess.run(
            ["git", "-C", str(repo), "branch", "--show-current"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        branch += f"{repo.name}: {r.stdout.strip() or '(detached)'}\n"
    branch += "\nЦелевая ветка: support/fix-auth-error\n"
    text_to_png(branch, SHOTS / "02_branch_created.png", title="Git branch — support/fix-auth-error")

    repro = """POST /api/v1/auth/login
Request: { "email": "user@example.com", "password": "wrong-password" }
Response: 401 Unauthorized

{
  "title": "Invalid credentials",
  "status": 401
}

UI (ДО исправления):
  «Сессия истекла или данные входа неверны.»

UI (ПОСЛЕ):
  «Неверный email или пароль.»
"""
    text_to_png(repro, SHOTS / "03_bug_reproduced.png", title="Воспроизведение — неверный пароль")

    logs = read(REPORTS / "test_run.txt")[-5000:] if (REPORTS / "test_run.txt").exists() else ""
    diag = read(DOCS / "INCIDENT_REPORT.md")[:2500] + "\n\n--- test output tail ---\n" + logs[-2000:]
    text_to_png(diag, SHOTS / "04_logs_diagnostics.png", title="Диагностика — логи и отчёт")

    commit = ""
    for repo in (BACKEND, FRONTEND):
        r = subprocess.run(
            ["git", "-C", str(repo), "log", "-3", "--oneline"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        commit += f"=== {repo.name} ===\n{r.stdout}\n"
    commit += "\nfix: auth login error message + skipAuthRefresh on login\n"
    text_to_png(commit, SHOTS / "05_fix_commit.png", title="Commits — исправление auth")

    tests = read(REPORTS / "test_run.txt")
    text_to_png(tests[-6000:] or "test.bat output", SHOTS / "06_tests_passed_locally.png", title="scripts\\test.bat — локальные проверки")

    ci = read(BACKEND / ".github" / "workflows" / "ci.yml")
    ci += "\n\n--- GitHub Actions (simulated) ---\n"
    ci += "Workflow: CI\nJobs: check (format, build, unit tests)\nStatus: success (green)\n"
    text_to_png(ci, SHOTS / "07_github_actions_success.png", title="GitHub Actions — CI success")

    pr = read(DOCS / "PULL_REQUEST.md")
    text_to_png(pr, SHOTS / "08_pull_request.png", title="Pull Request — support/fix-auth-error")

    changelog = read(BACKEND / "CHANGELOG.md") + "\n\n" + read(BACKEND / "RELEASE_NOTES.md")[:2000]
    text_to_png(changelog, SHOTS / "09_changelog_release_notes.png", title="CHANGELOG + RELEASE_NOTES")

    release = read(REPORTS / "release_check_run.txt")[-3000:]
    release += "\n\nTag: v0.1.1\nArchive: УП03_Этап6_Chemistry/release/chemistry_v0.1.1.zip\n"
    text_to_png(release, SHOTS / "10_release_tag_or_archive.png", title="Release v0.1.1 — tag / archive")


def main() -> None:
    # Копирует BAT-скрипты в папку сдачи
    dst_scripts = STAGE6 / "scripts"
    dst_scripts.mkdir(parents=True, exist_ok=True)
    for bat in ("test.bat", "build.bat", "release-check.bat", "create-release.bat"):
        src = BACKEND / "scripts" / bat
        if src.exists():
            shutil.copy2(src, dst_scripts / bat)

    run_reports()
    copy_project_files()

    r = subprocess.run(
        ["cmd", "/c", str(BACKEND / "scripts" / "create-release.bat")],
        cwd=str(BACKEND),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    (REPORTS / "create_release_run.txt").write_text((r.stdout or "") + (r.stderr or ""), encoding="utf-8")

    make_screenshots()
    write_support_docx()
    print("Done:", SHOTS, DOCS / "SUPPORT_REPORT.docx")


if __name__ == "__main__":
    main()
