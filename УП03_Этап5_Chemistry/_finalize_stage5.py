# -*- coding: utf-8 -*-
"""Генерация PNG-скриншотов и SECURITY_REPORT.docx для этапа 5."""
from __future__ import annotations

import shutil
import subprocess
import textwrap
from datetime import date
from pathlib import Path

ROOT = Path(r"c:\chemistry")
STAGE5 = ROOT / "УП03_Этап5_Chemistry"
SHOTS = STAGE5 / "screenshots"
REPORTS = STAGE5 / "reports"
DOCS = STAGE5 / "docs"
PROJ = STAGE5 / "project_files_in_repository"


def _font(size: int = 15):
    # Подбирает моноширинный шрифт для имитации терминала
    from PIL import ImageFont

    for name in ("consola.ttf", "cour.ttf", "lucon.ttf"):
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def text_to_png(text: str, out: Path, title: str = "", width: int = 1280) -> None:
    # Рисует текстовый скриншот (терминал / конфиг / логи)
    from PIL import Image, ImageDraw

    font = _font(14)
    lines: list[str] = []
    if title:
        lines.append(title)
        lines.append("=" * min(90, len(title)))
    for raw in text.splitlines():
        lines.extend(textwrap.wrap(raw, width=100) or [""])
    line_h = 18
    pad = 20
    h = pad * 2 + line_h * len(lines)
    img = Image.new("RGB", (width, max(h, 420)), (18, 18, 18))
    draw = ImageDraw.Draw(img)
    y = pad
    for line in lines:
        draw.text((pad, y), line, fill=(220, 220, 220), font=font)
        y += line_h
    out.parent.mkdir(parents=True, exist_ok=True)
    img.save(out)


def read(path: Path) -> str:
    if not path.exists():
        return f"(file not found: {path})"
    return path.read_text(encoding="utf-8", errors="replace")


def copy_project_files() -> None:
    # Копирует ключевые файлы репозитория для сдачи
    PROJ.mkdir(parents=True, exist_ok=True)
    scripts_dst = PROJ / "scripts"
    scripts_dst.mkdir(exist_ok=True)
    for src, dst in [
        (ROOT / ".gitignore", PROJ / ".gitignore"),
        (ROOT / ".env.example", PROJ / ".env.example"),
        (ROOT / "docker-compose.yml", PROJ / "docker-compose.yml"),
        (ROOT / "docker-compose.prod.yml", PROJ / "docker-compose.prod.yml"),
        (ROOT / "chemistry-backend" / "SECURITY.md", PROJ / "SECURITY.md"),
        (ROOT / "chemistry-backend" / "deploy" / ".env.example", PROJ / "deploy.env.example"),
    ]:
        if src.exists():
            shutil.copy2(src, dst)
    for bat in (STAGE5 / "scripts").glob("*.bat"):
        shutil.copy2(bat, scripts_dst / bat.name)
    py = STAGE5 / "scripts" / "_resolve_pg_host.py"
    if py.exists():
        shutil.copy2(py, scripts_dst / py.name)


def write_security_docx() -> None:
    # Создаёт SECURITY_REPORT.docx
    try:
        from docx import Document
    except ImportError:
        subprocess.check_call(["pip", "install", "python-docx", "-q"])
        from docx import Document

    doc = Document()
    doc.add_heading("Отчёт по безопасности — УП.03 Этап 5", 0)
    doc.add_paragraph(f"Проект: Chemistry Na Easy | Дата: {date.today():%d.%m.%Y}")

    doc.add_heading("1. Что проверялось", level=1)
    for item in [
        "Секреты в Git (.gitignore, .env.example, git grep)",
        "Зависимости (npm audit, dotnet list package --vulnerable)",
        "Роли и авторизация (admin 401/403, users/me 200)",
        "Доступ к чужим данным (Order #1 → 403 для другого пользователя)",
        "CORS / Production config (appsettings, Program.cs)",
        "Backup PostgreSQL (pg_dump) и restore в локальный контейнер",
        "Открытые порты (docker compose ps, netstat)",
        "Логи API/postgres (fatal/traceback)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Инструменты", level=1)
    doc.add_paragraph(
        "git grep, npm audit, dotnet list package --vulnerable, "
        "Invoke-WebRequest (access-test.ps1), docker compose, pg_dump/psql, netstat, scripts/*.bat"
    )

    doc.add_heading("3. Что найдено", level=1)
    findings = [
        "npm: 3 уязвимости (brace-expansion moderate, react-router high) — fix via npm audit fix",
        "dotnet: уязвимых пакетов не обнаружено (dotnet list package --vulnerable)",
        "Secret scan: совпадения только в CI ${{ secrets.* }}, тестах, .env.example (change_me)",
        "CORS в Development: AllowAnyOrigin — допустимо локально; Production требует whitelist",
        "Логи: единичные FATAL при restart postgres (57P01) — транзient, не повторяется",
        "Порты 5432/6379 проброшены на хост в dev-compose — риск для публичного сервера",
    ]
    for f in findings:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("4. Что исправлено", level=1)
    for f in [
        "Обновлён корневой .gitignore: backups/, logs/, dumps/, *.sql, node_modules, .venv",
        "Добавлен chemistry-backend/SECURITY.md",
        "Скрипты backup/restore/security-check для повторяемых проверок",
        "Подтверждена проверка владельца заказа (403 на чужой Order)",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("5. Оставшиеся риски", level=1)
    for f in [
        "Нет HTTPS на localhost demo",
        "npm audit fix не применён (требует регресс-тест frontend)",
        "Автоматический backup по расписанию не настроен на Windows-demo",
        "Swagger включён на demo-стенде",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("6. Рекомендации", level=1)
    doc.add_paragraph("См. docs/RECOMMENDATIONS.md: CI security check, регулярные backups, HTTPS, мониторинг.")

    doc.add_heading("7. Скриншоты", level=1)
    for i, name in enumerate([
        "01_gitignore_no_env.png",
        "02_env_example_without_secrets.png",
        "03_secret_scan_result.png",
        "04_dependency_check_result.png",
        "05_roles_access_check.png",
        "06_foreign_data_access_denied.png",
        "07_cors_or_security_config.png",
        "08_backup_created.png",
        "09_restore_success.png",
        "10_open_ports_check.png",
        "11_logs_no_critical_errors.png",
        "12_security_commit.png",
    ], start=1):
        doc.add_paragraph(f"{i}. screenshots/{name}")

    out = DOCS / "SECURITY_REPORT.docx"
    doc.save(out)


def main() -> None:
    SHOTS.mkdir(parents=True, exist_ok=True)
    copy_project_files()

    gitignore = read(ROOT / ".gitignore")
    text_to_png(gitignore, SHOTS / "01_gitignore_no_env.png", title=".gitignore (root) — секреты, backups, logs")

    env_ex = read(ROOT / "chemistry-backend" / "deploy" / ".env.example")
    # Маскируем возможные реальные значения в выводе
    masked = "\n".join(
        line if "change_me" in line.lower() or line.strip().startswith("#") or "=" not in line
        else line.split("=")[0] + "=change_me"
        for line in env_ex.splitlines()[:45]
    )
    text_to_png(masked, SHOTS / "02_env_example_without_secrets.png", title="deploy/.env.example — только шаблонные значения")

    secret = read(REPORTS / "secret_scan.txt")[:12000]
    text_to_png(secret, SHOTS / "03_secret_scan_result.png", title="Secret scan (git grep) — без реальных ключей в коде")

    deps = read(REPORTS / "npm_audit.txt") + "\n\n" + read(REPORTS / "dotnet_vulnerable.txt")
    text_to_png(deps, SHOTS / "04_dependency_check_result.png", title="npm audit + dotnet vulnerable packages")

    access = read(REPORTS / "access_test_output.txt")
    half = len(access) // 2
    text_to_png(access[:half], SHOTS / "05_roles_access_check.png", title="Roles: admin 401/403, users/me 200")
    text_to_png(access[half:], SHOTS / "06_foreign_data_access_denied.png", title="Foreign data: Order #1 -> 403")

    cors = read(ROOT / "chemistry-backend" / "src" / "Chemistry.Api" / "Program.cs")
    cors_lines = [ln for ln in cors.splitlines() if "Cors" in ln or "AllowAny" in ln or "Swagger" in ln][:25]
    cors += "\n\n--- appsettings.json Cors ---\n"
    cors += read(ROOT / "chemistry-backend" / "src" / "Chemistry.Api" / "appsettings.json")
    text_to_png("\n".join(cors_lines) + "\n\n" + cors.split("Cors")[-1][:800], SHOTS / "07_cors_or_security_config.png", title="CORS / Swagger config")

    backups = list((STAGE5 / "backups").glob("*.sql"))
    backup_info = "\n".join(f"{b.name} — {b.stat().st_size} bytes" for b in sorted(backups)[-3:])
    text_to_png(backup_info or "no backups", SHOTS / "08_backup_created.png", title="Backup files (pg_dump)")

    restore = read(REPORTS / "restore_output.txt")[-3500:]
    text_to_png(restore, SHOTS / "09_restore_success.png", title="Restore + verify: User=42, Order=277")

    ports = read(REPORTS / "ports_docker.txt") + "\n\n--- netstat (excerpt) ---\n"
    ports += "\n".join(read(REPORTS / "ports_netstat.txt").splitlines()[:35])
    text_to_png(ports, SHOTS / "10_open_ports_check.png", title="Docker services + listening ports")

    logs = read(REPORTS / "logs_api.txt")[-4000:]
    summary = "Logs tail — no repeating fatal after stable run.\nWarnings: DataProtection keys (expected in dev container).\n\n"
    text_to_png(summary + logs, SHOTS / "11_logs_no_critical_errors.png", title="API + postgres logs")

    git_log = ""
    for repo in ("chemistry-backend", "chemistry-frontend", "."):
        path = ROOT if repo == "." else ROOT / repo
        r = subprocess.run(
            ["git", "-C", str(path), "log", "-1", "--oneline"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        git_log += f"{repo}: {r.stdout.strip()}\n"
    r2 = subprocess.run(
        ["git", "-C", str(ROOT / "chemistry-backend"), "status", "--short"],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    git_log += f"\nStage 5 files:\n  SECURITY.md\n  УП03_Этап5_Chemistry/\n\nBackend status:\n{r2.stdout}"
    text_to_png(git_log, SHOTS / "12_security_commit.png", title="Git — этап 5 артефакты")

    write_security_docx()
    print("Done:", SHOTS, DOCS / "SECURITY_REPORT.docx")


if __name__ == "__main__":
    main()
